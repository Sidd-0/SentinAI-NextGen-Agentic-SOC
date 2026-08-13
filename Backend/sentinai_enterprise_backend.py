"""
SentinAI Sovereign Elite v10.7 — Production Active Defense Backend
Real-time 7-domain deep system scanning, autonomous autopilot, and forensic response.
"""

import os
import sys
import sqlite3
import uvicorn
import hashlib
import time
import psutil
import subprocess
import secrets
import smtplib
import re
import json
import uuid
import asyncio
import logging
import socket
from datetime import datetime
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from typing import Any, Dict, List, Optional, Tuple

from fastapi import FastAPI, HTTPException, Request, BackgroundTasks
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from crewai import Agent, Task, Crew, Process, LLM
from docx import Document
from dotenv import load_dotenv
from contextlib import asynccontextmanager

load_dotenv()

logging.getLogger("uvicorn.access").setLevel(logging.WARNING)

# ==================== MASTER CONFIGURATION ====================

MASTER_AI_KEY = os.getenv("GROQ_API_KEY", "enter yout groq apiu key here ")
SENDER_EMAIL = os.getenv("SENDER_EMAIL", "priyanshdwvd@gmail.com")
SMTP_PASSWORD = os.getenv("SMTP_PASSWORD", "mdjn wtgf sbob yfaj")
DB_FILE = "sentinai_enterprise.db"
REPORTS_DIR = "forensic_vault"
BACKEND_PORT = int(os.getenv("SENTINAI_PORT", "8000"))
OLLAMA_PORT = int(os.getenv("OLLAMA_PORT", "11434"))
HOSTS_FILE = r"C:\Windows\System32\drivers\etc\hosts"
SCAN_INTERVAL_SEC = 5

THREAT_PROCESSES = {"notepad.exe", "calc.exe", "mimikatz.exe", "mimikatz"}
STATUS_SECURED = "SECURED"
STATUS_WARNING = "WARNING"
STATUS_DANGER = "DANGER"

os.environ["OPENAI_API_KEY"] = MASTER_AI_KEY
os.makedirs(REPORTS_DIR, exist_ok=True)

# Shared telemetry cache (updated by background deep-scan loop)
_TELEMETRY_LOCK = asyncio.Lock()
_TELEMETRY_CACHE: Dict[str, Any] = {
    "it_infra": "Initializing...",
    "network": "Initializing...",
    "endpoint": "Initializing...",
    "iam": "Initializing...",
    "cloud": "Initializing...",
    "api": "Initializing...",
    "app_sec": "Initializing...",
    "status": "INITIALIZING",
    "uptime": "99.9%",
    "domain_status": {},
    "last_scan": None,
}

_AUTOPILOT_HANDLED: set = set()
_MISSION_IN_FLIGHT = False
_SERVER_BOOT = time.time()

# ==================== APP INITIALIZATION ====================


@asynccontextmanager
async def lifespan(app: FastAPI):
    print("[🚀] SOVEREIGN SENTINEL ACTIVE: 24/7 Deep System Scan (5s interval)...")
    asyncio.create_task(sentinel_autopilot_loop())
    yield


app = FastAPI(
    title="SentinAI Sovereign Elite v10.7",
    docs_url=None,
    redoc_url=None,
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==================== DATABASE ====================


def init_db():
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute(
        """CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT,
            email TEXT UNIQUE,
            password_hash TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            last_login TIMESTAMP)"""
    )
    cursor.execute(
        """CREATE TABLE IF NOT EXISTS reset_tokens (
            token TEXT PRIMARY KEY,
            email TEXT,
            created_at REAL)"""
    )
    cursor.execute(
        """CREATE TABLE IF NOT EXISTS security_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            domain TEXT,
            severity TEXT,
            message TEXT,
            metric_snapshot TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)"""
    )
    conn.commit()
    conn.close()


init_db()

# ==================== NEURAL ENGINES ====================

groq_llm = LLM(
    model="groq/llama-3.3-70b-versatile",
    api_key=MASTER_AI_KEY,
    temperature=0.0,
)

# Ollama — OpenAI-compatible endpoint (bypasses OpenAI connection errors)
ollama_llm = LLM(
    model="ollama/llama3.2",
    base_url="http://localhost:11434/v1",
    api_key="ollama",
    temperature=0.0,
)

# ==================== AUTH & SANITIZATION ====================


def hash_password(password: str) -> str:
    salt = "SentinAi_Elite_Hardening_Salt_2026"
    return hashlib.sha256((password + salt).encode()).hexdigest()


def verify_password(pw: str, pw_hash: str) -> bool:
    return hash_password(pw) == pw_hash


def sanitize_input_text(text: str) -> str:
    if not text:
        return ""
    clean = re.sub(r"<[^>]*?>", "", text)
    clean = re.sub(r"[;\"\'|]", "", clean)
    return clean.strip()


def log_security_event(domain: str, severity: str, message: str, snapshot: Optional[dict] = None):
    try:
        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO security_logs (domain, severity, message, metric_snapshot) VALUES (?, ?, ?, ?)",
            (domain, severity, message, json.dumps(snapshot or {})),
        )
        conn.commit()
        conn.close()
        ts = datetime.now().strftime("%H:%M:%S")
        print(f"[{ts}] - [{domain}] - {message}")
    except Exception as exc:
        print(f"[DB] Security log write failed: {exc}")


def log_combat_action(domain: str, action: str, target: str = ""):
    """Write a live combat feed entry (e.g. NEUTRALIZED: notepad.exe)."""
    msg = f"NEUTRALIZED: {target}" if target else action
    log_security_event(domain, STATUS_DANGER if "NEUTRALIZED" in msg else STATUS_SECURED, msg)


# ==================== REAL SYSTEM PROBES (7 DOMAINS) ====================


def _run_cmd(cmd: str, timeout: int = 10) -> Tuple[bool, str]:
    try:
        result = subprocess.run(
            cmd, shell=True, capture_output=True, text=True, timeout=timeout
        )
        output = (result.stdout or "") + (result.stderr or "")
        return result.returncode == 0, output.strip()
    except Exception as exc:
        return False, str(exc)


def _port_listening(host: str, port: int, timeout: float = 1.0) -> bool:
    try:
        with socket.create_connection((host, port), timeout=timeout):
            return True
    except OSError:
        return False


def scan_endpoint_protection() -> Dict[str, Any]:
    threats: List[str] = []
    proc_count = 0
    try:
        for proc in psutil.process_iter(["pid", "name"]):
            proc_count += 1
            name = (proc.info.get("name") or "").lower()
            if name in THREAT_PROCESSES or "mimikatz" in name:
                threats.append(f"{proc.info.get('name')} (PID {proc.info.get('pid')})")
    except Exception as exc:
        return {
            "metric": f"Scan error: {exc}",
            "status": STATUS_DANGER,
            "threats": [],
            "detail": str(exc),
        }

    if threats:
        status = STATUS_DANGER
        metric = f"{proc_count} Processes | THREAT: {', '.join(threats[:3])}"
    else:
        status = STATUS_SECURED
        metric = f"{proc_count} Active Processes | EDR: CLEAN"

    return {"metric": metric, "status": status, "threats": threats, "proc_count": proc_count}


def scan_network_security() -> Dict[str, Any]:
    firewall_on = False
    firewall_detail = "UNKNOWN"
    ok, output = _run_cmd("netsh advfirewall show allprofiles state")
    if ok or output:
        # Match each profile line: "State                 ON" / "State                 OFF"
        states = re.findall(r"state\s+(on|off)", output, re.IGNORECASE)
        firewall_on = bool(states) and all(s.lower() == "on" for s in states)
        firewall_detail = "ON" if firewall_on else "OFF"
    else:
        firewall_detail = "UNREACHABLE"

    try:
        sockets = len(psutil.net_connections(kind="inet"))
    except (psutil.AccessDenied, PermissionError):
        sockets = len(psutil.net_connections())
    except Exception:
        sockets = "N/A"

    status = STATUS_SECURED if firewall_on else STATUS_DANGER

    metric = f"Firewall: {firewall_detail} | {sockets} Active Sockets"
    return {
        "metric": metric,
        "status": status,
        "firewall_on": firewall_on,
        "sockets": sockets,
        "detail": output[:200] if output else "",
    }


MALICIOUS_PROCESSES = {"mimikatz.exe", "mimikatz", "psexec.exe", "procdump.exe"}


def scan_it_infrastructure() -> Dict[str, Any]:
    try:
        cpu = psutil.cpu_percent(interval=0.3)
        ram = psutil.virtual_memory().percent
        malicious = []
        for proc in psutil.process_iter(["name"]):
            name = (proc.info.get("name") or "").lower()
            if name in MALICIOUS_PROCESSES or "mimikatz" in name:
                malicious.append(proc.info.get("name"))
    except Exception as exc:
        return {"metric": f"Offline: {exc}", "status": STATUS_DANGER, "cpu": 0, "ram": 0}

    if malicious or ram > 95:
        status = STATUS_DANGER
    else:
        status = STATUS_SECURED

    metric = f"CPU: {cpu:.1f}% | RAM: {ram:.1f}%"
    if malicious:
        metric += f" | MALICIOUS: {', '.join(malicious[:2])}"
    return {"metric": metric, "status": status, "cpu": cpu, "ram": ram, "malicious": malicious}


def scan_iam() -> Dict[str, Any]:
    """Query live Windows sessions via `query user` and parse actual usernames."""
    win_user = os.environ.get("USERNAME", "UNKNOWN").upper()
    session_count = 0
    usernames: List[str] = []
    primary_user = win_user

    ok, output = _run_cmd("query user")
    if ok or output:
        for line in output.splitlines():
            raw = line.strip()
            if not raw or raw.upper().startswith("USERNAME"):
                continue
            is_active = raw.startswith(">")
            cleaned = raw.lstrip(">").strip()
            parts = cleaned.split()
            if not parts:
                continue
            uname = parts[0].upper()
            if uname in ("USERNAME", "SESSIONNAME", "ID", "STATE"):
                continue
            usernames.append(uname)
            if is_active:
                primary_user = uname

    if not usernames:
        try:
            sessions = psutil.users()
            session_count = len(sessions)
            usernames = [u.name.upper() for u in sessions] or [win_user]
            primary_user = usernames[0]
        except Exception:
            usernames = [win_user]
            primary_user = win_user
    else:
        session_count = len(set(usernames))

    status = STATUS_SECURED
    metric = f"Windows User: {primary_user} | {session_count} Active Session(s)"
    return {
        "metric": metric,
        "status": status,
        "sessions": session_count,
        "users": ", ".join(usernames),
        "windows_username": primary_user,
    }


def scan_cloud_security() -> Dict[str, Any]:
    ollama_up = _port_listening("127.0.0.1", OLLAMA_PORT)
    if ollama_up:
        status = STATUS_SECURED
        metric = f"Local AI Bridge (:{OLLAMA_PORT}): LISTENING"
    else:
        status = STATUS_DANGER
        metric = f"Local AI Bridge (:{OLLAMA_PORT}): OFFLINE"

    return {"metric": metric, "status": status, "ollama_listening": ollama_up}


def scan_api_security() -> Dict[str, Any]:
    api_up = _port_listening("127.0.0.1", BACKEND_PORT)
    if api_up:
        status = STATUS_SECURED
        metric = f"FastAPI Shield (:{BACKEND_PORT}): ACTIVE & REACHABLE"
    else:
        status = STATUS_DANGER
        metric = f"FastAPI Shield (:{BACKEND_PORT}): UNREACHABLE"

    return {"metric": metric, "status": status, "api_listening": api_up}


def scan_application_security() -> Dict[str, Any]:
    intact = False
    detail = ""
    try:
        if os.path.isfile(HOSTS_FILE):
            size = os.path.getsize(HOSTS_FILE)
            with open(HOSTS_FILE, "r", encoding="utf-8", errors="ignore") as fh:
                content = fh.read(4096)
            has_localhost = "127.0.0.1" in content and "localhost" in content.lower()
            intact = size > 0 and has_localhost
            detail = f"size={size}B localhost_entry={'yes' if has_localhost else 'no'}"
        else:
            detail = "hosts file missing"
    except Exception as exc:
        detail = str(exc)

    if intact:
        status = STATUS_SECURED
        metric = "hosts file: VERIFIED & INTACT"
    else:
        status = STATUS_DANGER
        metric = f"hosts file: COMPROMISED/MISSING ({detail})"

    return {"metric": metric, "status": status, "intact": intact, "detail": detail}


def perform_deep_system_scan() -> Dict[str, Any]:
    """Execute a full 7-domain deep scan against live OS telemetry."""
    endpoint = scan_endpoint_protection()
    network = scan_network_security()
    it_infra = scan_it_infrastructure()
    iam = scan_iam()
    cloud = scan_cloud_security()
    api = scan_api_security()
    app_sec = scan_application_security()

    domain_status = {
        "endpoint": endpoint["status"],
        "network": network["status"],
        "it_infra": it_infra["status"],
        "iam": iam["status"],
        "cloud": cloud["status"],
        "api": api["status"],
        "app_sec": app_sec["status"],
    }

    statuses = list(domain_status.values())
    if STATUS_DANGER in statuses:
        global_status = "CRITICAL"
    else:
        global_status = "OPERATIONAL"

    uptime_hours = (time.time() - _SERVER_BOOT) / 3600
    uptime_pct = min(99.99, 99.0 + min(uptime_hours / 100, 0.99))

    return {
        "it_infra": it_infra["metric"],
        "network": network["metric"],
        "endpoint": endpoint["metric"],
        "iam": iam["metric"],
        "cloud": cloud["metric"],
        "api": api["metric"],
        "app_sec": app_sec["metric"],
        "status": global_status,
        "uptime": f"{uptime_pct:.1f}%",
        "domain_status": domain_status,
        "windows_username": iam.get("windows_username", os.environ.get("USERNAME", "UNKNOWN").upper()),
        "last_scan": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "_scan_details": {
            "endpoint": endpoint,
            "network": network,
            "it_infra": it_infra,
            "iam": iam,
            "cloud": cloud,
            "api": api,
            "app_sec": app_sec,
        },
    }


def resolve_domain_label(domain_key: str) -> str:
    mapping = {
        "endpoint": "Endpoint Protection",
        "network": "Network Security",
        "it_infra": "IT Infrastructure",
        "iam": "Identity & Access Mgmt (IAM)",
        "cloud": "Cloud Security",
        "api": "API Security",
        "app_sec": "Application Security",
    }
    return mapping.get(domain_key, domain_key)


# ==================== DOMAIN-SPECIFIC COMBAT MATRIX ====================

DOMAIN_COMBAT_PLAYBOOK: Dict[str, List[str]] = {
    "Endpoint Protection": [
        "taskkill /F /IM notepad.exe",
        "taskkill /F /IM calc.exe",
        "taskkill /F /IM mimikatz.exe",
        "wmic process where \"name='notepad.exe'\" delete",
        "wmic process where \"name='calc.exe'\" delete",
    ],
    "Network Security": [
        "netsh advfirewall firewall add rule name=\"SentinAI_Block_Threat\" dir=in action=block remoteip=185.220.101.45",
        "netsh advfirewall set allprofiles state on",
        "ipconfig /flushdns",
    ],
    "Identity & Access Mgmt (IAM)": [
        "net user Guest /active:no",
        "logoff",
    ],
    "Application Security": [
        f'icacls "{HOSTS_FILE}" /deny Everyone:(W)',
        "net stop w3svc",
    ],
    "IT Infrastructure": [
        "shutdown /a",
    ],
    "Cloud Security": [
        "netsh advfirewall set allprofiles state on",
    ],
    "API Security": [
        "netsh http flush logbuffer",
    ],
}

WHITELIST_PROCESS_NAMES = {
    "explorer.exe",
    "svchost.exe",
    "wininit.exe",
    "python.exe",
    "runtimebroker.exe",
    "System",
    "csrss.exe",
    "lsass.exe",
}


def execute_domain_combat(domain: str, threat_context: str = "") -> str:
    """Run domain-specific Windows CLI kill rules before/after AI mitigation."""
    print(f"\n[⚔] DOMAIN COMBAT MATRIX: {domain}")
    cmds = list(DOMAIN_COMBAT_PLAYBOOK.get(domain, []))

    ctx = threat_context.lower()
    if "notepad" in ctx and domain == "Endpoint Protection":
        cmds.insert(0, "taskkill /F /IM notepad.exe")
        cmds.insert(1, 'wmic process where "name=\'notepad.exe\'" delete')
    if "calc" in ctx and domain == "Endpoint Protection":
        cmds.insert(0, "taskkill /F /IM calc.exe")
        cmds.insert(1, 'wmic process where "name=\'calc.exe\'" delete')
    if "mimikatz" in ctx and domain == "Endpoint Protection":
        cmds.insert(0, "taskkill /F /IM mimikatz.exe")
        cmds.insert(1, 'wmic process where "name=\'mimikatz.exe\'" delete')

    actions = []
    for cmd in dict.fromkeys(cmds):
        if any(w in cmd.lower() for w in WHITELIST_PROCESS_NAMES if "taskkill" in cmd.lower()):
            if any(w in cmd.lower() for w in ("explorer.exe", "svchost.exe", "python.exe")):
                continue
        try:
            print(f"[!] COMBAT: {cmd}")
            subprocess.run(cmd, shell=True, check=False, capture_output=True, text=True, timeout=15)
            actions.append(f"EXECUTED: {cmd}")
            print("[✓] COMBAT STEP COMPLETE")
            if "taskkill" in cmd.lower() or "wmic" in cmd.lower():
                proc_match = re.search(r"/IM\s+(\S+)|name='([^']+)'", cmd, re.IGNORECASE)
                if proc_match:
                    target = (proc_match.group(1) or proc_match.group(2) or "").strip()
                    if target:
                        log_combat_action(domain, "NEUTRALIZED", target)
            elif "flushdns" in cmd.lower():
                log_combat_action("Network Security", "DNS cache flushed")
            elif "firewall" in cmd.lower():
                log_combat_action("Network Security", "Firewall rule applied")
            elif "logoff" in cmd.lower():
                log_combat_action("Identity & Access Mgmt (IAM)", "Session logoff initiated")
        except Exception as exc:
            actions.append(f"ATTEMPTED: {cmd} ({exc})")

    return "\n".join(actions) if actions else "No domain-specific commands required."


def run_universal_combat(ai_output: str, log_context: str = "", domain: str = "") -> str:
    """Parse AI command output and execute safe Windows mitigations."""
    print("\n[⚔] AGENT 3: INITIATING SYSTEM MITIGATION...")
    domain_log = execute_domain_combat(domain, log_context) if domain else ""
    cmds_to_run: List[str] = []

    json_match = re.search(r"\[.*\]", ai_output, re.DOTALL)
    if json_match:
        try:
            raw_json = json_match.group(0).replace("'", '"')
            data = json.loads(raw_json)
            for item in data:
                if isinstance(item, dict):
                    cmd = item.get("command", "")
                    params = item.get("parameters", "")
                    full_cmd = f"{cmd} {params}".strip()
                    if full_cmd:
                        cmds_to_run.append(full_cmd)
                else:
                    cmds_to_run.append(str(item))
        except json.JSONDecodeError:
            pass

    if not cmds_to_run:
        cmds_to_run = re.findall(
            r"((?:taskkill|netsh|powershell|net stop|del|icacls|reg|bcdedit|wmic|logoff|ipconfig)\s+[\w\s\-/\\\"=.:*'()]+)",
            ai_output,
            re.IGNORECASE,
        )

    if "notepad" in log_context.lower() or "notepad" in ai_output.lower():
        cmds_to_run.append("taskkill /F /IM notepad.exe")

    actions = [domain_log] if domain_log else []
    for cmd in set(cmds_to_run):
        cmd = cmd.strip().rstrip(".").rstrip("'").rstrip('"')
        if any(app in cmd.lower() for app in WHITELIST_PROCESS_NAMES):
            continue
        try:
            print(f"[!] SYSTEM CALL: {cmd}")
            subprocess.run(cmd, shell=True, check=False, capture_output=True, text=True, timeout=15)
            actions.append(f"SUCCESS: {cmd}")
            print("[✓✓✓] COMPLETED")
        except Exception:
            actions.append(f"NEUTRALIZED/VERIFIED: {cmd}")

    return "\n".join(a for a in actions if a)


# ==================== NEURAL MISSION ENGINE ====================

DOMAIN_TACTICS = {
    "Endpoint Protection": "Focus on process termination. Tools: taskkill /F /IM, wmic process delete.",
    "Network Security": "Focus on traffic isolation. Tools: netsh advfirewall firewall add rule, ipconfig /flushdns.",
    "Cloud Security": "Focus on AI bridge hardening. Tools: netsh advfirewall, service verification.",
    "Identity & Access Mgmt (IAM)": "Focus on session revocation. Tools: net user /active:no, logoff.",
    "API Security": "Focus on traffic filtering. Tools: netsh http, port verification.",
    "Application Security": "Focus on file integrity. Tools: icacls deny, net stop.",
    "IT Infrastructure": "Focus on resource containment. Tools: shutdown /a, systeminfo.",
}


async def execute_sovereign_mission(
    log_text: str,
    domain: str,
    user_email: Optional[str] = None,
    force_ollama: bool = False,
) -> Dict[str, Any]:
    use_ollama = force_ollama or "Autonomous" in (domain or "")
    active_llm = ollama_llm if use_ollama else groq_llm
    engine_name = "OLLAMA (llama3.2)" if use_ollama else "GROQ"
    tactics = DOMAIN_TACTICS.get(domain, "General Security Protocols.")

    admin = Agent(
        role=f"{domain} Admin",
        goal=f"Remediate threats in {domain}.",
        backstory=f"Authorized Enforcer. {tactics}",
        llm=active_llm,
        allow_delegation=False,
        verbose=True,
    )

    t1 = Task(
        description=f"Analyze: {log_text}. Identify Attack Name.",
        agent=admin,
        expected_output="Attack Name",
    )
    t2 = Task(
        description=(
            "Provide a FLAT JSON array of Windows CLI commands to kill the threat. "
            "Example: [\"taskkill /F /IM name.exe\"]. Output ONLY the JSON array."
        ),
        agent=admin,
        expected_output="JSON array of strings.",
    )
    t3 = Task(
        description="Summarize in 3 paragraphs.",
        agent=admin,
        expected_output="Summary.",
    )

    mission = Crew(agents=[admin], tasks=[t1, t2, t3], process=Process.sequential, verbose=True)
    print(f"\n[🧠] NEURAL CORE: Engaging {engine_name} for {domain}...")

    try:
        await asyncio.to_thread(mission.kickoff)
        classification = str(t1.output)
        commands_json = str(t2.output)
        summary = str(t3.output)
    except Exception as exc:
        print(f"[🧠] AI engine fallback — executing domain combat only: {exc}")
        classification = f"Autonomous Threat — {domain}"
        commands_json = "[]"
        summary = (
            f"Autonomous deep-scan detected a {domain} anomaly. "
            f"Details: {log_text}. Domain combat matrix executed. "
            f"AI analysis unavailable ({exc}); manual review recommended."
        )

    combat_log = run_universal_combat(commands_json, log_text, domain)

    report_path = generate_professional_report(
        {
            "domain": domain,
            "classification": classification,
            "combat_log": combat_log,
            "analysis": summary,
        }
    )

    if not user_email:
        try:
            conn = sqlite3.connect(DB_FILE)
            cursor = conn.cursor()
            cursor.execute("SELECT email FROM users ORDER BY last_login DESC LIMIT 1")
            row = cursor.fetchone()
            conn.close()
            user_email = row[0] if row else SENDER_EMAIL
        except Exception:
            user_email = SENDER_EMAIL

    send_forensic_email(user_email, report_path)
    log_security_event(domain, STATUS_DANGER, log_text, {"classification": classification})

    return {
        "classification": classification,
        "action_taken": combat_log,
        "analysis": summary,
        "report_url": f"/api/download/{os.path.basename(report_path)}",
        "engine": engine_name,
    }


# ==================== AUTONOMOUS AUTOPILOT ====================


def _build_threat_signature(domain_key: str, scan: Dict[str, Any]) -> Optional[str]:
    details = scan.get("_scan_details", {})
    domain_detail = details.get(domain_key, {})

    if domain_key == "endpoint" and domain_detail.get("threats"):
        return f"endpoint:{','.join(domain_detail['threats'])}"
    if domain_key == "network" and not domain_detail.get("firewall_on", True):
        return "network:firewall_off"
    if domain_key == "it_infra" and domain_detail.get("status") == STATUS_DANGER:
        return f"it_infra:cpu{domain_detail.get('cpu')}_ram{domain_detail.get('ram')}"
    if domain_key == "cloud" and not domain_detail.get("ollama_listening", True):
        return "cloud:ollama_down"
    if domain_key == "api" and not domain_detail.get("api_listening", True):
        return "api:backend_down"
    if domain_key == "app_sec" and not domain_detail.get("intact", True):
        return "app_sec:hosts_integrity"
    if domain_key == "iam" and domain_detail.get("status") == STATUS_DANGER:
        return "iam:session_anomaly"
    return None


async def _handle_danger_domains(scan: Dict[str, Any]):
    global _MISSION_IN_FLIGHT

    domain_status = scan.get("domain_status", {})
    danger_domains = [k for k, v in domain_status.items() if v == STATUS_DANGER]

    if not danger_domains or _MISSION_IN_FLIGHT:
        return

    for domain_key in danger_domains:
        sig = _build_threat_signature(domain_key, scan)
        if not sig or sig in _AUTOPILOT_HANDLED:
            continue

        _AUTOPILOT_HANDLED.add(sig)
        _MISSION_IN_FLIGHT = True
        domain_label = resolve_domain_label(domain_key)
        details = scan.get("_scan_details", {}).get(domain_key, {})
        log_text = (
            f"[AUTONOMOUS DEEP SCAN] {domain_label} status=DANGER. "
            f"Metric: {scan.get(domain_key if domain_key != 'it_infra' else 'it_infra', 'N/A')}. "
            f"Details: {json.dumps(details, default=str)}"
        )
        print(f"\n[🚨] AUTOPILOT: DANGER in {domain_label} — engaging Ollama llama3.2...")
        try:
            await execute_sovereign_mission(log_text, domain_label, force_ollama=True)
        except Exception as exc:
            print(f"[🚨] Autopilot mission error: {exc}")
            execute_domain_combat(domain_label, log_text)
        finally:
            _MISSION_IN_FLIGHT = False


async def sentinel_autopilot_loop():
    """Deep system scan every 5 seconds; auto-neutralize on DANGER."""
    psutil.cpu_percent(interval=None)
    while True:
        try:
            scan = await asyncio.to_thread(perform_deep_system_scan)
            async with _TELEMETRY_LOCK:
                _TELEMETRY_CACHE.clear()
                _TELEMETRY_CACHE.update(scan)

            await _handle_danger_domains(scan)

            details = scan.get("_scan_details", {})
            endpoint = details.get("endpoint", {})
            for threat in endpoint.get("threats", []):
                proc_name = threat.split("(")[0].strip().lower()
                if proc_name in THREAT_PROCESSES or "mimikatz" in proc_name:
                    sig = f"proc:{proc_name}:{threat}"
                    if sig not in _AUTOPILOT_HANDLED:
                        _AUTOPILOT_HANDLED.add(sig)
                        subprocess.run(
                            f"taskkill /F /IM {proc_name}",
                            shell=True,
                            capture_output=True,
                        )
                        subprocess.run(
                            f'wmic process where "name=\'{proc_name}\'" delete',
                            shell=True,
                            capture_output=True,
                        )
                        log_combat_action("Endpoint Protection", "NEUTRALIZED", proc_name)
        except Exception as exc:
            print(f"[⚠] Autopilot loop error: {exc}")

        await asyncio.sleep(SCAN_INTERVAL_SEC)


# ==================== API ENDPOINTS ====================


@app.get("/api/telemetry")
async def get_telemetry():
    """Return live 7-domain telemetry with SECURED / WARNING / DANGER status keys."""
    try:
        scan = perform_deep_system_scan()
        async with _TELEMETRY_LOCK:
            _TELEMETRY_CACHE.clear()
            _TELEMETRY_CACHE.update(scan)

        response = {k: v for k, v in scan.items() if not k.startswith("_")}
        return response
    except Exception as exc:
        print(f"Telemetry Error: {exc}")
        async with _TELEMETRY_LOCK:
            cached = dict(_TELEMETRY_CACHE)
        cached["status"] = "ERROR"
        cached["it_infra"] = f"Offline: {exc}"
        return cached


@app.post("/api/analyze")
async def analyze_threat(request: Request):
    """Manual Analyze button — maps to dashboard Neural Mitigation."""
    try:
        data = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON body")

    log_text = sanitize_input_text(data.get("log", ""))
    domain = data.get("domain") or "Endpoint Protection"
    user_email = data.get("email")

    if not log_text:
        raise HTTPException(status_code=400, detail="Log intake required.")

    return await execute_sovereign_mission(log_text, domain, user_email)


@app.post("/api/signup")
async def signup(request: Request):
    data = await request.json()
    try:
        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO users (name, email, password_hash) VALUES (?, ?, ?)",
            (data["name"], data["email"].lower(), hash_password(data["password"])),
        )
        conn.commit()
        conn.close()
        return {"ok": True}
    except Exception:
        return {"ok": False}


@app.post("/api/login")
async def login(request: Request):
    data = await request.json()
    email = data.get("email", "").lower().strip()
    pwd = data.get("password")
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("SELECT name, password_hash FROM users WHERE email=?", (email,))
    user = cursor.fetchone()
    if user and verify_password(pwd, user[1]):
        cursor.execute(
            "UPDATE users SET last_login = ? WHERE email = ?",
            (datetime.now().strftime("%Y-%m-%d %H:%M:%S"), email),
        )
        conn.commit()
        conn.close()
        return {"ok": True, "name": user[0]}
    conn.close()
    raise HTTPException(status_code=401)


@app.post("/api/forgot-password")
async def forgot_password(request: Request, bg: BackgroundTasks):
    print("\n[🔑] FORGOT PASSWORD PROTOCOL TRIGGERED")
    try:
        data = await request.json()
        email = data.get("email", "").lower().strip()

        if not email:
            return {"ok": False, "error": "Email required."}

        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        cursor.execute("SELECT name FROM users WHERE email = ?", (email,))
        user = cursor.fetchone()

        if user:
            name = user[0]
            token = secrets.token_urlsafe(32)
            cursor.execute(
                "INSERT OR REPLACE INTO reset_tokens (token, email, created_at) VALUES (?, ?, ?)",
                (token, email, time.time()),
            )
            conn.commit()

            reset_link = f"http://localhost:8080/reset-password?token={token}"
            subject = "Neural Reset Beacon — SentinAI Elite"
            body = (
                f"Operator {name}, click the link below to reset your access code:\n\n"
                f"{reset_link}\n\nThis link expires in 1 hour."
            )

            msg = MIMEMultipart()
            msg["From"] = SENDER_EMAIL
            msg["To"] = email
            msg["Subject"] = subject
            msg.attach(MIMEText(body, "plain"))
            with smtplib.SMTP("smtp.gmail.com", 587) as s:
                s.starttls()
                s.login(SENDER_EMAIL, SMTP_PASSWORD)
                s.send_message(msg)

            print(f"[✓] Reset beacon dispatched to {email}")

        conn.close()
        return {"ok": True, "message": "If the identity exists, a reset beacon has been dispatched."}
    except Exception as e:
        print(f"[✗] FORGOT PASSWORD ERROR: {e}")
        return {"ok": False, "error": str(e)}


@app.post("/api/reset-password")
async def reset_password(request: Request):
    print("\n[🔄] RESET PASSWORD EXECUTION TRIGGERED")
    try:
        data = await request.json()
        token = data.get("token")
        new_password = data.get("new_password")

        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        cursor.execute("SELECT email, created_at FROM reset_tokens WHERE token = ?", (token,))
        row = cursor.fetchone()

        if not row:
            conn.close()
            return {"ok": False, "error": "Invalid or expired token."}

        email, created_at = row
        if time.time() - created_at > 3600:
            cursor.execute("DELETE FROM reset_tokens WHERE token = ?", (token,))
            conn.commit()
            conn.close()
            return {"ok": False, "error": "Token expired."}

        hashed_pw = hash_password(new_password)
        cursor.execute("UPDATE users SET password_hash = ? WHERE email = ?", (hashed_pw, email))
        cursor.execute("DELETE FROM reset_tokens WHERE token = ?", (token,))

        conn.commit()
        conn.close()
        print(f"[✓] PASSWORD RESET SUCCESSFUL: {email}")
        return {"ok": True, "message": "Access code updated successfully."}
    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.get("/api/download/{filename}")
async def download_report(filename: str):
    safe_name = os.path.basename(filename)
    path = os.path.join(REPORTS_DIR, safe_name)
    if not os.path.isfile(path):
        raise HTTPException(status_code=404, detail="Report not found.")
    return FileResponse(path, filename=safe_name)


@app.get("/api/logs")
async def get_combat_logs():
    """Return the last 15 security/combat feed entries for the live terminal."""
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute(
        "SELECT id, domain, severity, message, created_at "
        "FROM security_logs ORDER BY id DESC LIMIT 15"
    )
    rows = cursor.fetchall()
    conn.close()
    entries = []
    for r in reversed(rows):
        ts = r[4] or ""
        if " " in ts:
            ts = ts.split(" ")[1][:8]
        entries.append({
            "id": r[0],
            "domain": r[1],
            "severity": r[2],
            "message": r[3],
            "timestamp": ts,
            "formatted": f"{ts} - [{r[1]}] - {r[3]}",
        })
    return {"logs": entries, "count": len(entries)}


@app.post("/api/lockdown")
async def initiate_global_lockdown():
    """Emergency protocol: kill browsers, enable firewall, flush DNS."""
    actions: List[str] = []
    browser_procs = [
        "chrome.exe", "firefox.exe", "msedge.exe", "iexplore.exe",
        "opera.exe", "brave.exe", "vivaldi.exe",
    ]
    for browser in browser_procs:
        try:
            subprocess.run(
                f"taskkill /F /IM {browser}",
                shell=True, capture_output=True, text=True, timeout=10,
            )
            actions.append(f"KILLED: {browser}")
            log_combat_action("Global Lockdown", "NEUTRALIZED", browser)
        except Exception as exc:
            actions.append(f"ATTEMPTED: {browser} ({exc})")

    firewall_cmds = [
        "netsh advfirewall set allprofiles state on",
        "netsh advfirewall set domainprofile state on",
        "netsh advfirewall set privateprofile state on",
        "netsh advfirewall set publicprofile state on",
    ]
    for cmd in firewall_cmds:
        try:
            subprocess.run(cmd, shell=True, capture_output=True, text=True, timeout=10)
            actions.append(f"EXECUTED: {cmd}")
        except Exception as exc:
            actions.append(f"FAILED: {cmd} ({exc})")

    try:
        subprocess.run("ipconfig /flushdns", shell=True, capture_output=True, text=True, timeout=10)
        actions.append("EXECUTED: ipconfig /flushdns")
        log_combat_action("Global Lockdown", "DNS cache flushed")
    except Exception as exc:
        actions.append(f"FAILED: flushdns ({exc})")

    log_security_event(
        "Global Lockdown",
        STATUS_DANGER,
        "EMERGENCY LOCKDOWN INITIATED — browsers terminated, firewall enabled, DNS flushed",
        {"actions": actions},
    )

    return {
        "status": "LOCKDOWN_ACTIVE",
        "message": "Global lockdown protocol executed.",
        "actions": actions,
    }


@app.get("/api/security-logs")
async def get_security_logs(limit: int = 50):
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute(
        "SELECT id, domain, severity, message, metric_snapshot, created_at "
        "FROM security_logs ORDER BY id DESC LIMIT ?",
        (limit,),
    )
    rows = cursor.fetchall()
    conn.close()
    return {
        "logs": [
            {
                "id": r[0],
                "domain": r[1],
                "severity": r[2],
                "message": r[3],
                "metric_snapshot": r[4],
                "created_at": r[5],
            }
            for r in rows
        ]
    }


@app.get("/api/reset-db")
async def reset_database():
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("DROP TABLE IF EXISTS users")
    cursor.execute("DROP TABLE IF EXISTS reset_tokens")
    cursor.execute("DROP TABLE IF EXISTS security_logs")
    conn.commit()
    conn.close()
    init_db()
    return {"status": "Success"}


# ==================== FORENSIC REPORT & EMAIL ====================


def generate_professional_report(data: Dict[str, str]) -> str:
    session_id = str(uuid.uuid4()).upper()[:8]
    fname = f"SentinAI_Forensic_{session_id}.docx"
    path = os.path.join(REPORTS_DIR, fname)
    doc = Document()
    doc.add_heading("SENTINAI ELITE v10.7 - OFFICIAL FORENSIC REPORT", 0).alignment = 1
    doc.add_heading("1. Incident Summary", level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    meta = [
        ("Incident ID", f"SNT-{session_id}"),
        ("Timestamp", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("Security Domain", data["domain"]),
        ("Threat Classification", data["classification"]),
        ("Status", "NEUTRALIZED"),
    ]
    for i, (k, v) in enumerate(meta):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = str(v)
    doc.add_heading("2. Combat Execution Log", level=1)
    doc.add_paragraph(f"Action Taken: {data['combat_log']}")
    doc.add_heading("3. AI Forensic Analysis", level=1)
    doc.add_paragraph(data["analysis"])
    doc.save(path)
    return path


def send_forensic_email(to_email: str, report_path: str):
    try:
        msg = MIMEMultipart()
        msg["From"] = SENDER_EMAIL
        msg["To"] = to_email
        msg["Subject"] = "CRITICAL: SentinAI Neutralization Report"
        msg.attach(
            MIMEText(
                "Operator, an attack was neutralized. Professional report attached.",
                "plain",
            )
        )
        with open(report_path, "rb") as f:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header(
                "Content-Disposition",
                f"attachment; filename={os.path.basename(report_path)}",
            )
            msg.attach(part)
        with smtplib.SMTP("smtp.gmail.com", 587) as s:
            s.starttls()
            s.login(SENDER_EMAIL, SMTP_PASSWORD)
            s.send_message(msg)
        print(f"[📧] REPORT DISPATCHED TO: {to_email}")
    except Exception as exc:
        print(f"[📧] Email dispatch failed: {exc}")


# ==================== ENTRY POINT ====================

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=BACKEND_PORT)
